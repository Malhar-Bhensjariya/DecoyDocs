#!/usr/bin/env python3
"""
simple_generate.py
Lightweight Gemini-based document generator for DecoyDocs backend.
Generates a single DOCX without similarity checks (no sentence-transformers dependency).
"""

import os
import sys
import json
import uuid
from pathlib import Path
from datetime import datetime, timezone

from dotenv import load_dotenv
from google import genai
import docx
from docx.shared import Inches
import markdown2
from bs4 import BeautifulSoup

# Load API key
dotenv_path = Path(__file__).resolve().parent / ".env"
if dotenv_path.exists():
    load_dotenv(dotenv_path)

api_key = os.getenv("GEMINI_API_KEY")
if not api_key:
    raise ValueError("Missing GEMINI_API_KEY in environment or .env file")

client = genai.Client(api_key=api_key)

TEMPLATE_PROMPTS = {
    "generic_report": "Generate a professional generic business report about {title}. Include executive summary, key findings, and recommendations.",
    "employee_bonus": "Generate an HR document about {title} regarding employee bonuses. Include department breakdown, bonus distribution, and payment schedule.",
    "q3_financial": "Generate a financial report titled {title} for Q3. Include revenue, expenses, profit margins, and trend analysis.",
    "hr_review": "Generate an HR review document titled {title}. Include performance metrics, feedback, and development plans.",
    "sales_pipeline": "Generate a sales document about {title}. Include sales opportunities, pipeline status, and forecast.",
}

def generate_docx_from_gemini(title: str, template: str, output_path: str):
    """Generate DOCX using Gemini content."""
    prompt = TEMPLATE_PROMPTS.get(template, f"Write professional content about {title}.")
    prompt = prompt.format(title=title)
    
    # Call Gemini
    response = client.models.generate_content(
        model="gemini-2.5-flash",
        contents=prompt
    )
    
    content = response.text
    
    # Convert markdown to HTML then to DOCX
    html = markdown2.markdown(content, extras=['fenced-code-blocks', 'tables'])
    
    # Create DOCX
    doc = docx.Document()
    doc.add_heading(title, level=1)
    
    # Parse HTML and add to doc
    soup = BeautifulSoup(html, "html.parser")
    for elem in soup.find_all(['p', 'h1', 'h2', 'h3', 'ul', 'ol', 'blockquote']):
        if elem.name in ('h1', 'h2', 'h3'):
            level = int(elem.name[1])
            doc.add_heading(elem.get_text(), level=level)
        elif elem.name == 'p':
            doc.add_paragraph(elem.get_text())
        elif elem.name == 'ul':
            for li in elem.find_all('li', recursive=False):
                doc.add_paragraph(li.get_text(), style='List Bullet')
        elif elem.name == 'ol':
            for li in elem.find_all('li', recursive=False):
                doc.add_paragraph(li.get_text(), style='List Number')
        elif elem.name == 'blockquote':
            p = doc.add_paragraph(elem.get_text())
            p.style = 'Intense Quote'
    
    doc.save(output_path)

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: simple_generate.py <title> <template> [output_dir]")
        sys.exit(1)
    
    title = sys.argv[1]
    template = sys.argv[2]
    output_dir = sys.argv[3] if len(sys.argv) > 3 else "generated_docs"
    
    os.makedirs(output_dir, exist_ok=True)
    
    # Create clean filename
    clean_title = "".join(c for c in title if c.isalnum() or c in ' -_')[:40]
    docx_name = f"{clean_title}_{str(uuid.uuid4())[:8]}.docx"
    output_path = os.path.join(output_dir, docx_name)
    
    try:
        generate_docx_from_gemini(title, template, output_path)
        print(json.dumps({"success": True, "path": output_path, "filename": docx_name}))
    except Exception as e:
        print(json.dumps({"success": False, "error": str(e)}))
        sys.exit(1)
