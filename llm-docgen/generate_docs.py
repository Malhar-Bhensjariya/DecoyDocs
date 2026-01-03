#!/usr/bin/env python3
"""
generate_docs.py
Phase 1: Generate realistic document .docx files using Google Gemini.
Creates professional-looking sample docs (HR, business, etc.) using a local templates.py file.

This version:
- Loads GEMINI_API_KEY from llm-docgen/.env using python-dotenv
- Converts common Markdown produced by the model into Word formatting
- Does NOT expose generation metadata in the visible document (keeps metadata in core properties and JSON sidecar)
- Produces shorter, clean filenames: <clean-title>_<short-uuid>.docx
"""

import os
import re
import sys
import argparse
import json
import uuid
import time
from datetime import datetime, timezone
from pathlib import Path
from typing import Optional, List

from dotenv import load_dotenv
from google import genai
import docx
from docx.shared import Inches
import markdown2
from bs4 import BeautifulSoup

# Import Templates
try:
    from templates import TEMPLATES, SAMPLE_TITLES
except Exception:
    print("Missing templates.py or import error. Ensure it exists in the same folder.")
    raise

DEFAULT_MODEL = "gemini-2.5-pro"


def load_api_key_from_env_file() -> Optional[str]:
    """Load GEMINI_API_KEY from .env and set GOOGLE_API_KEY for compatibility."""
    dotenv_path = Path(__file__).resolve().parent / ".env"
    if dotenv_path.exists():
        load_dotenv(dotenv_path)
    key = os.getenv("GEMINI_API_KEY") or os.getenv("GOOGLE_API_KEY")
    if key:
        os.environ["GOOGLE_API_KEY"] = key
    return key


def init_client(api_key: Optional[str]) -> genai.Client:
    """Initialize the Gemini client with API key or fallback."""
    try:
        return genai.Client(api_key=api_key) if api_key else genai.Client()
    except TypeError:
        if api_key:
            os.environ["GOOGLE_API_KEY"] = api_key
        return genai.Client()

def clean_text(text: str) -> str:
    import re
    text = text.strip()
    lines = text.splitlines()

    # Remove duplicate title or markdown heading
    if len(lines) > 1 and lines[0].strip('# ').lower() in lines[1].strip('# ').lower():
        lines = lines[1:]
    lines = [line.lstrip('#').strip() for line in lines]
    text = '\n'.join(lines)

    # Remove orphan bullets
    text = re.sub(r'\n•\s*\n', '\n', text)
    text = re.sub(r'•\s*\n', '', text)
    text = re.sub(r'•\s{0,2}', '', text)

    # Remove stray final bullet
    text = re.sub(r'•\s*$', '', text)
    return text


def generate_text(client, prompt: str, model: str, max_output_tokens: int) -> str:
    """Generate text using Gemini."""
    resp = client.models.generate_content(model=model, contents=prompt)
    return getattr(resp, "text", str(resp))


def ensure_output_folder(script_file: str) -> Path:
    """Create or return the generated_docs folder one level above script dir."""
    output_dir = Path(script_file).resolve().parent.parent / "generated_docs"
    output_dir.mkdir(parents=True, exist_ok=True)
    return output_dir


def sanitize_filename(name: str) -> str:
    """Remove invalid chars and truncate for safe filenames."""
    keep = "-_.() abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ0123456789"
    name = "".join(c for c in name if c in keep).strip().replace(" ", "_")
    return name[:60]


def short_uuid() -> str:
    """Return a short 8-character UUID."""
    return str(uuid.uuid4())[:8]


def _html_to_docx_elements(doc: docx.Document, html: str):
    """
    Convert an HTML fragment (from markdown2) into docx paragraphs, headings, lists, and tables.
    """
    soup = BeautifulSoup(html, "html.parser")

    for elem in soup.body.children if soup.body else soup.children:
        if getattr(elem, "name", None) in ("h1", "h2", "h3", "h4", "h5", "h6"):
            level = int(elem.name[1]) if elem.name[1].isdigit() else 2
            # python-docx supports levels 0-9, keep reasonable mapping
            level = min(max(level, 1), 4)
            doc.add_heading(elem.get_text().strip(), level=level)
        elif getattr(elem, "name", None) == "p":
            # handle inline bold/italic/links by creating runs
            p = doc.add_paragraph()
            _add_runs_from_html(p, elem)
        elif getattr(elem, "name", None) in ("ul", "ol"):
            for li in elem.find_all("li", recursive=False):
                if elem.name == "ul":
                    doc.add_paragraph(li.get_text().strip(), style="List Bullet")
                else:
                    doc.add_paragraph(li.get_text().strip(), style="List Number")
        elif getattr(elem, "name", None) == "blockquote":
            p = doc.add_paragraph()
            p.style = "Intense Quote" if "Intense Quote" in [s.name for s in doc.styles] else p.style
            _add_runs_from_html(p, elem)
        elif getattr(elem, "name", None) == "pre":
            # code block
            p = doc.add_paragraph(elem.get_text())
            p.style = p.style
        elif getattr(elem, "name", None) == "table":
            # build a table
            rows = elem.find_all("tr")
            if not rows:
                continue
            ncols = len(rows[0].find_all(["td", "th"]))
            table = doc.add_table(rows=0, cols=ncols)
            table.style = "Table Grid" if "Table Grid" in [s.name for s in doc.styles] else table.style
            for r in rows:
                cells = r.find_all(["td", "th"])
                row_cells = table.add_row().cells
                for idx, cell in enumerate(cells):
                    row_cells[idx].text = cell.get_text().strip()
        else:
            # fallback: text node or unknown tag
            text = getattr(elem, "string", None)
            if text and text.strip():
                doc.add_paragraph(text.strip())


def _add_runs_from_html(paragraph: docx.text.paragraph.Paragraph, elem):
    """
    Add runs to a paragraph handling <strong>, <em>, <a> tags.
    """
    for node in elem.children:
        if getattr(node, "name", None) == "strong" or getattr(node, "name", None) == "b":
            run = paragraph.add_run(node.get_text())
            run.bold = True
        elif getattr(node, "name", None) == "em" or getattr(node, "name", None) == "i":
            run = paragraph.add_run(node.get_text())
            run.italic = True
        elif getattr(node, "name", None) == "a":
            # Represent links as text (link target in parentheses)
            text = node.get_text().strip()
            href = node.get("href", "").strip()
            run = paragraph.add_run(f"{text} ({href})" if href else text)
        elif getattr(node, "name", None) == "code":
            run = paragraph.add_run(node.get_text())
            # optional: style monospace by setting font.name if available
            try:
                run.font.name = "Courier New"
            except Exception:
                pass
        else:
            # text node or nested tags
            s = node.string if getattr(node, "string", None) is not None else None
            if s and s.strip():
                paragraph.add_run(s.strip())
            elif getattr(node, "name", None):
                # nested structure, recurse
                _add_runs_from_html(paragraph, node)


def create_docx(title: str, body: str, metadata: dict, output_path: Path):
    """
    Create a professional-looking .docx file from LLM-generated text.
    Ensures realistic formatting (proper spacing, black headings, consistent fonts, and real bullet lists).
    """
    # Set default font for LibreOffice compatibility
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = docx.Document()
    # Set default font to LibreOffice compatible
    style = doc.styles['Normal']
    style.font.name = 'Liberation Serif'
    style = doc.styles['Normal']
    style.font.name = 'Liberation Serif'  # LibreOffice default
    style.font.size = Pt(11)

    # ---------------------------
    # Title — large, bold, centered, LibreOffice compatible font
    # ---------------------------
    title_paragraph = doc.add_paragraph()
    title_run = title_paragraph.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(20)
    title_run.font.name = 'Liberation Serif'  # LibreOffice compatible
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Small space below title
    title_paragraph.paragraph_format.space_after = Pt(10)

    # ---------------------------
    # Split LLM text into paragraphs
    # ---------------------------
    paragraphs = [p.strip() for p in body.split("\n") if p.strip()]
    paragraphs = [p for p in paragraphs if not re.match(r"^[•*\-]\s*$", p)]
    paragraphs = [re.sub(r"\*\*(.*?)\*\*", r"\1", p) for p in paragraphs]

    for p_text in paragraphs:
        # 1️⃣ Detect bullet list items (start with '*', '-', or '•')
        if p_text.lstrip().startswith(("*", "-", "•")):
            clean = (
                p_text.lstrip("*•- \t")  # remove bullet symbol and spaces
                .replace("**", "")       # remove markdown bold markers
                .strip()
            )
            para = doc.add_paragraph(style="List Bullet")
            run = para.add_run(clean)
            run.font.size = Pt(11)
            run.font.name = "Liberation Serif"
            run.font.color.rgb = RGBColor(0, 0, 0)
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after = Pt(2)

        # 2️⃣ Detect metadata lines (TO, FROM, DATE, SUBJECT)
        elif any(p_text.startswith(x) for x in ["TO:", "FROM:", "DATE:", "SUBJECT:"]):
            para = doc.add_paragraph()
            run = para.add_run(p_text)
            run.bold = True
            run.font.size = Pt(11)
            run.font.name = "Liberation Serif"
            run.font.color.rgb = RGBColor(0, 0, 0)
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(2)

        # 3️⃣ Detect headings (uppercase, ends with ':', starts with ### or **)
        elif (
            p_text.isupper()
            or p_text.endswith(":")
            or p_text.startswith("###")
            or p_text.startswith("**")
        ):
            clean = p_text.replace("*", "").replace("#", "").strip()
            para = doc.add_paragraph()
            run = para.add_run(clean)
            run.bold = True
            run.font.size = Pt(14)
            run.font.name = "Liberation Serif"
            run.font.color.rgb = RGBColor(0, 0, 0)
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            # tighter vertical spacing (was causing too much gap earlier)
            para.paragraph_format.space_before = Pt(8)
            para.paragraph_format.space_after = Pt(4)

        # 4️⃣ Normal body paragraph
        else:
            para = doc.add_paragraph(p_text)
            para.style = "Normal"
            for run in para.runs:
                run.font.size = Pt(11)
                run.font.name = "Liberation Serif"  # LibreOffice compatible
                run.font.color.rgb = RGBColor(0, 0, 0)
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(6)

    # ---------------------------
    # Metadata stored in file properties (not visible)
    # ---------------------------
    try:
        core = doc.core_properties
        core.title = title
        core.comments = (
            f"uuid:{metadata['uuid']} | model:{metadata['model']} | "
            f"template:{metadata['template_used']} | generated_at:{metadata['generated_at']}"
        )
    except Exception:
        pass

    # ---------------------------
    # Save .docx and sidecar metadata JSON
    # ---------------------------
    filename = f"{sanitize_filename(title)}_{metadata['uuid'][:8]}.docx"
    docx_file = output_path / filename
    doc.save(str(docx_file))

    meta_file = output_path / f"{metadata['uuid']}.json"
    with open(meta_file, "w", encoding="utf-8") as f:
        json.dump(metadata, f, indent=2)

    return docx_file, meta_file



def extract_avoid_lists(text: str) -> tuple:
    """Extract topics and terminology to avoid in regeneration."""
    # Simple extraction: headings as topics, key nouns as terms
    lines = text.split('\n')
    topics = []
    terms = []
    for line in lines:
        line = line.strip()
        if line.isupper() or line.startswith('**') or line.startswith('###'):
            topics.append(line.replace('**', '').replace('#', '').strip())
        # Extract potential terms (simple: words >4 chars)
        words = [w.strip('.,') for w in line.split() if len(w) > 4 and w.isalpha()]
        terms.extend(words[:5])  # limit
    return list(set(topics)), list(set(terms))


def build_prompt(template_name: Optional[str], title: Optional[str], context: Optional[str], avoid_topics: list = None, avoid_terms: list = None) -> str:
    template_entry = TEMPLATES.get(template_name, TEMPLATES["generic_report"])
    prompt_text = template_entry["prompt"] if isinstance(template_entry, dict) else template_entry
    title_part = f"Document title: {title}\n\n" if title else ""
    context_part = f"Context:\n{context}\n\n" if context else ""

    avoid_part = ""
    if avoid_topics or avoid_terms:
        avoid_part = "\n\nAVOID the following to ensure uniqueness:\n"
        if avoid_topics:
            avoid_part += f"- Topics: {', '.join(avoid_topics)}\n"
        if avoid_terms:
            avoid_part += f"- Terminology: {', '.join(avoid_terms)}\n"
        avoid_part += "Make this document distinct in content, style, and focus.\n\n"

    return (
        f"{title_part}{context_part}{avoid_part}{prompt_text}\n\n"
        "Produce a realistic, structured professional document (HR, business, finance, etc.). "
        "Use headings and short paragraphs. Prefer plain text or Markdown (use headings, bold, lists, and simple tables). "
        "Do not include notes about being AI-generated or include internal debug metadata in the visible document."
    )



def main(argv: List[str]):
    parser = argparse.ArgumentParser(description="LLM Document Generator (Gemini)")
    parser.add_argument("--count", type=int, default=3, help="Number of documents to generate (default 3)")
    parser.add_argument("--template", type=str, default=None)
    parser.add_argument("--title", type=str, default=None)
    parser.add_argument("--context", type=str, default=None)
    parser.add_argument("--model", type=str, default=DEFAULT_MODEL)
    parser.add_argument("--max-tokens", type=int, default=800)
    parser.add_argument("--similarity-threshold", type=float, default=0.80)
    args = parser.parse_args(argv)

    api_key = load_api_key_from_env_file()
    if api_key:
        print("Using GEMINI_API_KEY from .env")
    else:
        print("No API key found; using default auth (ADC or GOOGLE_API_KEY env var if set)")

    client = init_client(api_key)
    out_dir = ensure_output_folder(__file__)
    print(f"Output folder: {out_dir}")

    # Import similarity module
    sys.path.append(str(Path(__file__).resolve().parent.parent))
    from similarity import check_similarity_threshold, compute_similarity_matrix

    generated_docs = []  # list of (uuid, text, title, metadata)

    for i in range(args.count):
        template_data = TEMPLATES.get(args.template, TEMPLATES["generic_report"])
        base_title = args.title or (template_data["title"] if isinstance(template_data, dict) else "Company Report")
        uid = short_uuid()
        metadata = {
            "uuid": uid,
            "generated_at": datetime.now(timezone.utc).isoformat(),
            "model": args.model,
            "prompt_summary": base_title,
            "template_used": args.template or "generic_report",
            "source_script": Path(__file__).name,
        }

        # Build avoid lists from previous docs
        avoid_topics = []
        avoid_terms = []
        if generated_docs:
            for prev_uuid, prev_text, _, _ in generated_docs:
                t, terms = extract_avoid_lists(prev_text)
                avoid_topics.extend(t)
                avoid_terms.extend(terms)
            avoid_topics = list(set(avoid_topics))
            avoid_terms = list(set(avoid_terms))

        # Vary title, department, style for uniqueness
        variations = [
            ("", "", ""),
            (" - Marketing Department", "Focus on marketing metrics and campaigns.", "Use a promotional tone."),
            (" - Engineering Division", "Emphasize technical achievements and R&D.", "Adopt a technical writing style.")
        ]
        var_title, var_context, var_style = variations[i % len(variations)]
        title = base_title + var_title
        context = (args.context or "") + " " + var_context + " " + var_style

        prompt = build_prompt(args.template, title, context, avoid_topics, avoid_terms)

        print(f"Generating document {i+1}/{args.count}: '{title}' (uuid={uid}) ...")
        try:
            text = generate_text(client, prompt, model=args.model, max_output_tokens=args.max_tokens)
            text = clean_text(text)

            # Check similarity if not first doc
            if generated_docs:
                temp_docs = generated_docs + [(uid, text)]
                if check_similarity_threshold(temp_docs, args.similarity_threshold):
                    print(f"  Similarity too high, regenerating...")
                    # Regenerate with stronger avoid
                    avoid_topics.extend(["budget", "performance", "strategy"])  # add defaults
                    prompt = build_prompt(args.template, title, context, avoid_topics, avoid_terms)
                    text = generate_text(client, prompt, model=args.model, max_output_tokens=args.max_tokens)
                    text = clean_text(text)
                    # Check again
                    if check_similarity_threshold(generated_docs + [(uid, text)], args.similarity_threshold):
                        print(f"  Still too similar, skipping for now (implement loop later)")
                        continue

            generated_docs.append((uid, text, title, metadata))
            docx_file, meta_file = create_docx(title, text, metadata, out_dir)
            print(f"Saved: {docx_file.name}  |  Metadata: {meta_file.name}")

        except Exception as e:
            print(f"Generation error: {e}")
            metadata["error"] = str(e)
            continue

        time.sleep(0.5)

    # Output similarity matrix
    if len(generated_docs) > 1:
        print("\nCosine Similarity Matrix:")
        matrix = compute_similarity_matrix(generated_docs)
        uuids = [doc[0] for doc in generated_docs]
        print("UUIDs:", uuids)
        print(matrix)


if __name__ == "__main__":
    main(sys.argv[1:])
