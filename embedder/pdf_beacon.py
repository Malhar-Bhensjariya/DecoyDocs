# pdf_beacon.py
# Embed active beacon triggers in PDF files
from pathlib import Path
from typing import Optional

try:
    from pypdf import PdfWriter, PdfReader
    from pypdf.generic import DictionaryObject, NameObject, ArrayObject
    HAS_PYPDF = True
except ImportError:
    try:
        from PyPDF2 import PdfFileWriter, PdfFileReader
        from PyPDF2.generic import DictionaryObject, NameObject, ArrayObject
        HAS_PYPDF = True
        PYPDF2_MODE = True
    except ImportError:
        HAS_PYPDF = False
        PYPDF2_MODE = False


def embed_beacon_in_pdf(pdf_path: Path, beacon_url: str, output_path: Optional[Path] = None) -> Path:
    """
    Embed an invisible beacon trigger in a PDF file.
    
    Creates an invisible annotation/link that loads the beacon URL when the PDF is opened.
    Uses PDF JavaScript action or invisible link annotation.
    
    Args:
        pdf_path: Path to input PDF
        beacon_url: URL to trigger (beacon endpoint)
        output_path: Optional output path (defaults to overwriting input)
    
    Returns:
        Path to the modified PDF
    """
    if not HAS_PYPDF:
        raise ImportError(
            "pypdf or PyPDF2 is required for PDF beacon embedding. "
            "Install with: pip install pypdf"
        )
    
    if output_path is None:
        output_path = pdf_path
    
    try:
        # Try pypdf (newer library)
        reader = PdfReader(str(pdf_path))
        writer = PdfWriter()
        
        # Copy all pages
        for page in reader.pages:
            writer.add_page(page)
        
        # Method 1: Add invisible link annotation (works when clicked, most reliable)
        first_page = writer.pages[0]
        mediabox = first_page.mediabox
        height = float(mediabox.height)
        link_annotation = DictionaryObject({
            NameObject("/Type"): NameObject("/Annot"),
            NameObject("/Subtype"): NameObject("/Link"),
            NameObject("/Rect"): ArrayObject([0, height - 1, 1, height]),
            NameObject("/Border"): ArrayObject([0, 0, 0]),  # Invisible
            NameObject("/A"): DictionaryObject({
                NameObject("/S"): NameObject("/URI"),
                NameObject("/URI"): beacon_url
            }),
            NameObject("/F"): 4,  # Invisible
        })
        
        if "/Annots" not in first_page:
            first_page[NameObject("/Annots")] = ArrayObject()
        annot_ref = writer._add_object(link_annotation)
        first_page[NameObject("/Annots")].append(annot_ref)
        
        # Method 2: Add JavaScript action (BEST-EFFORT ONLY)
        # NOTE: PDF JavaScript does NOT auto-execute in Evince (Linux Document Viewer)
        # and is often disabled in other viewers for security. This is best-effort only.
        # The primary auto-trigger should be DOCX remote images.
        try:
            js_code = f"""
            // Auto-trigger beacon when PDF opens (if JavaScript enabled)
            // NOTE: This will NOT work in Evince or most secure PDF viewers
            try {{
                var req = new XMLHttpRequest();
                req.open("GET", "{beacon_url}", true);
                req.send();
            }} catch(e) {{
                try {{
                    fetch("{beacon_url}");
                }} catch(e2) {{
                    var img = new Image();
                    img.src = "{beacon_url}";
                }}
            }}
            """
            writer.add_js(js_code)
            print(f"   JavaScript added (best-effort - disabled in Evince)")
        except Exception as js_err:
            pass  # JavaScript not available, that's okay
        
        # Write output
        output_path_str = str(output_path)
        with open(output_path_str, "wb") as out_file:
            writer.write(out_file)
        
        print(f"✅ Beacon embedded in PDF: {beacon_url}")
        return output_path
        
    except Exception as e:
        # Fallback: Try simpler approach with just URI action
        print(f"⚠️ Primary method failed ({e}), trying fallback...")
        return _embed_beacon_simple(pdf_path, beacon_url, output_path)


def _embed_beacon_simple(pdf_path: Path, beacon_url: str, output_path: Path) -> Path:
    """Simpler fallback: Add invisible link annotation only."""
    try:
        reader = PdfReader(str(pdf_path))
        writer = PdfWriter()
        
        for i, page in enumerate(reader.pages):
            new_page = writer.add_page(page)
            
            # Add invisible link on first page only
            if i == 0:
                mediabox = new_page.mediabox
                height = float(mediabox.height)
                
                # Create invisible link annotation
                link_annotation = DictionaryObject({
                    NameObject("/Type"): NameObject("/Annot"),
                    NameObject("/Subtype"): NameObject("/Link"),
                    NameObject("/Rect"): ArrayObject([0, height - 1, 1, height]),
                    NameObject("/Border"): ArrayObject([0, 0, 0]),
                    NameObject("/A"): DictionaryObject({
                        NameObject("/S"): NameObject("/URI"),
                        NameObject("/URI"): beacon_url
                    }),
                    NameObject("/F"): 4  # Invisible
                })
                
                # Add to page annotations
                if "/Annots" not in new_page:
                    new_page[NameObject("/Annots")] = ArrayObject()
                
                annot_ref = writer._add_object(link_annotation)
                new_page[NameObject("/Annots")].append(annot_ref)
        
        output_path_str = str(output_path)
        with open(output_path_str, "wb") as out_file:
            writer.write(out_file)
        
        print(f"✅ Beacon embedded (simple method) in PDF")
        return output_path
        
    except Exception as e:
        raise RuntimeError(f"Failed to embed beacon in PDF: {e}")


def embed_beacon_in_docx(docx_path: Path, beacon_url: str, output_path: Optional[Path] = None) -> Path:
    """
    Embed active beacon trigger in DOCX using python-docx with remote image.
    
    Uses python-docx to add a 1x1pt hidden remote image that auto-loads on open.
    This is LibreOffice-compatible and does not corrupt DOCX structure.
    
    NOTE: Remote images may be blocked by security settings in some environments.
    This is the most reliable auto-trigger method for DOCX files.
    
    Args:
        docx_path: Path to input DOCX
        beacon_url: URL to trigger (loaded as remote image)
        output_path: Optional output path
    
    Returns:
        Path to modified DOCX
    """
    from docx import Document
    from docx.shared import Pt
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls, qn
    
    if output_path is None:
        output_path = docx_path
    
    # Load document using python-docx (ensures valid structure)
    doc = Document(str(docx_path))
    
    # Add hidden paragraph at end for remote image
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.0
    
    # Add relationship for remote image (TargetMode="External")
    part = doc.part
    rel = part.rels.get_or_add_relationship(
        'http://schemas.openxmlformats.org/officeDocument/2006/relationships/image',
        beacon_url,
        is_external=True
    )
    
    # Create inline drawing with remote image reference
    # Use python-docx's XML structure to ensure validity
    drawing_xml = f'''
    <w:drawing xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
               xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
               xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main"
               xmlns:pic="http://schemas.openxmlformats.org/drawingml/2006/picture"
               xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships">
        <wp:inline distT="0" distB="0" distL="0" distR="0">
            <wp:extent cx="127" cy="127"/>
            <wp:effectExtent l="0" t="0" r="0" b="0"/>
            <wp:docPr id="1" name="beacon"/>
            <wp:cNvGraphicFramePr/>
            <a:graphic>
                <a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">
                    <pic:pic>
                        <pic:nvPicPr>
                            <pic:cNvPr id="0" name="beacon"/>
                            <pic:cNvPicPr/>
                        </pic:nvPicPr>
                        <pic:blipFill>
                            <a:blip r:link="{rel.rId}"/>
                            <a:stretch>
                                <a:fillRect/>
                            </a:stretch>
                        </pic:blipFill>
                        <pic:spPr>
                            <a:xfrm>
                                <a:off x="0" y="0"/>
                                <a:ext cx="127" cy="127"/>
                            </a:xfrm>
                            <a:prstGeom prst="rect">
                                <a:avLst/>
                            </a:prstGeom>
                        </pic:spPr>
                    </pic:pic>
                </a:graphicData>
            </a:graphic>
        </wp:inline>
    </w:drawing>
    '''
    
    # Parse and insert drawing
    drawing = parse_xml(drawing_xml)
    run = para.add_run()
    run._element.append(drawing)
    
    # Make run invisible (white text, tiny size)
    run.font.color.rgb = None  # Will be set to white
    from docx.shared import RGBColor
    run.font.color.rgb = RGBColor(255, 255, 255)
    run.font.size = Pt(1)
    
    # Save document (python-docx ensures valid OOXML structure)
    doc.save(str(output_path))
    
    return output_path
